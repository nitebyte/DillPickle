import os
import openai
import hashlib
import ebooklib
import csv
import requests
import json
from io import BytesIO
from ebooklib import epub
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



# Load your API key from an environment variable or secret management service
openai.api_key = "sk-yourapikeyhere"
api_key = os.environ.get("sk-yourapikeyhere")
url = "https://api.openai.com/v1/images/generations"
globaldesc=""
#------------------------------------------------------------------------------------------------------------------------------------------------

def PR(prompt, system_content, temperature=0.5, token=12, pp=0.5, fp=0.5):
    messages = [{'role': 'system', 'content': system_content}, {'role': 'user', 'content': prompt}]
    response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages, temperature=temperature, max_tokens=token, presence_penalty=pp,frequency_penalty=fp)
    system_message = response.choices[0].get('message', {}).get('content', '').strip()
    prompt_message = prompt.strip()
    #print("System message:", system_content)
    #print("Prompt message:", prompt_message)
    #print("Generated response:", system_message)
    return system_message

def APD(text, file_name):
    with open(file_name, "a+") as f:
        f.write(text + "\n")

def Title(text, place):
    return text[6:]

def txt_to_docx(txt_file, docx_file,town):
    # Create a new Document
    document = Document()
    # Set the page size and margins
    section = document.sections[0]
    section.page_width = Inches(4)
    section.page_height = Inches(6)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

     # Add title page
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(town)
    title_run.font.size = Pt(18)
    title_run.bold = True
    #title_paragraph.add_run("\n\nYour Guide to History, Culture, and Fun")
    document.add_page_break()

    # Add copyright page
    copyright_paragraph = document.add_paragraph()
    copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_paragraph.add_run("Copyright 2023").font.size = Pt(8)

    # Set the position of the copyright notice to the bottom of the copyright page
    copyright_paragraph_format = copyright_paragraph.paragraph_format
    copyright_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_paragraph_format.space_after = Pt(0)
    document.add_page_break()

    # Read the input text file
    with open(txt_file, 'r') as f:
        txt_lines = f.readlines()
        txt_lines = [line for line in txt_lines if line.strip()]

    # Process the text file lines
    for line in txt_lines:
        if line.startswith('H1'):
            document.add_page_break()
            title = line[2:].strip()
            heading = document.add_heading(level=1)
            heading_run = heading.add_run(title)
            heading_run.font.size = Pt(18)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer sk-yourapikeyhere"
            }
            imgprompt = PR("I want you to write a DALL-E image generation prompt to generate an image related to " + title + " in the town of " + town, "You write image generation prompts for DALL-E from the given input request. For example, if you were asked to write a prompt for an image about the geography of Damascus, Virginia, you may output: Mountainous terrain with dense forests and trees, peaceful and serene, in the vicinity of Damascus, VA, USA. Shot on a Canon EOS R6 with a Canon RF 24-105mm f/4L IS USM Lens, 4K film still, natural lighting, vibrant colors, crisp details, and soft shadows.", 0.5, 2048)
            print("Generating Image For: " + title + " : " + town)
            data = {
                "prompt": imgprompt,
                "n": 1,
                "size": "1024x1024"
            }

            response = requests.post(url, headers=headers, data=json.dumps(data))

            if response.status_code == 200:
                result = response.json()
                # Download the image from a link
                image_url = result['data'][0]['url']
                response = requests.get(image_url)
                img = BytesIO(response.content)
                # Add a paragraph with an image
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center the paragraph
                run = paragraph.add_run()
                run.add_picture(img, width=Inches(3))  # adjust width as necessary

        elif line.startswith('H2'):
            title = line[2:].strip()
            heading = document.add_heading(level=2)
            heading_run = heading.add_run(title)
            heading_run.font.size = Pt(16)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        else:
            content = line.strip()
            paragraph = document.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the generated DOCX file
    document.save(docx_file)


def read_towns_csv(csv_file):
    towns = []

    with open(csv_file, newline='') as f:
        reader = csv.reader(f)
        for row in reader:
            town, state = row[0].split(', ')
            towns.append([town, state])

    return towns
#------------------------------------------------------------------------------------------------------------------------------------------------

towns = read_towns_csv('towns.csv')

nouse=input();

for town_name, state in towns:
    town = town_name# + ", " + state
    savefile = town_name + " " + state + ".txt"

    book = [
    # Chapter 1
    [
    "H11.0 Introduction",
    "H21.1 Purpose of the Company Diligence Report",
    ],
    # Chapter 2
    [
    "H12.0 Company Overview",
    "H22.1 Company history and background",
    "H22.2 Ownership and governance structure",
    "H22.3 Organizational chart",
    "H22.4 Company culture and values",
    ],
    # Chapter 3
    [
    "H13.0 Financial Due Diligence",
    "H23.1 Financial statements (balance sheet, income statement, cash flow statement)",
    "H23.2 Financial ratios analysis",
    "H23.3 Key performance indicators (KPIs)",
    "H23.4 Historical financial trends",
    "H23.5 Forecasts and projections",
    ],
    # Chapter 4
    [
    "H14.0 Legal Due Diligence",
    "H24.1 Corporate documents (e.g., articles of incorporation, bylaws)",
    "H24.2 Contracts (e.g., customer agreements, supplier agreements)",
    "H24.3 Litigation history",
    "H24.4 Intellectual property rights",
    "H24.5 Regulatory compliance",
    ],
    # Chapter 5
    [
    "H15.0 Commercial Due Diligence",
    "H25.1 Market analysis",
    "H25.2 Competitive landscape",
    "H25.3 Sales and marketing strategy",
    "H25.4 Customer base analysis",
    "H25.5 Supply chain analysis",
    ],
    # Chapter 6
    [
    "H16.0 Operational Due Diligence",
    "H26.1 Production processes",
    "H26.2 Capacity utilization",
    "H26.3 Quality control processes",
    "H26.4 Supply chain management",
    "H26.5 IT systems and infrastructure",
    ],
    # Chapter 7
    [
    "H17.0 Human Resources Due Diligence",
    "H27.1 Employee demographics",
    "H27.2 Compensation and benefits",
    "H27.3 Employee turnover rates",
    "H27.4 Training and development programs",
    "H27.5 Succession planning",
    ],
    # Chapter 8
    [
    "H18.0 Risk Assessment",
    "H28.1 Identification of risks",
    "H28.2 Evaluation of risks",
    "H28.3 Mitigation strategies",
    ],
    # Chapter 9
    [
    "H19.0 Conclusion",
    "H29.1 Summary of findings",
    "H29.2 Implications for potential investors",
    "H29.3 Recommendations",
    ]
    ]
    ShortBook = []

    for chapter in book:
        new_chapter = []
        for section in chapter:
            new_section = section[6:]
            new_chapter.append(new_section)
        ShortBook.append(new_chapter)

    total_chars = 0
    total_words = 0
    total_sections = sum([len(c) for c in book])
    completed_sections = 0

    print("Title: Company Diligence Report For "+town)

    for i, chapter in enumerate(book):
        for j, section in enumerate(chapter):
            print("Writing About The " + ShortBook[i][j] + " of " + town)
            if(j==0):
                prompt = "Write me a 1 paragraph brief introduction for a chapter about the " + ShortBook[i][j] + " of " + town + "."
            else:
                prompt = "Write me 8-10 extensive detailed and informative paragraphs about the " + ShortBook[i][j] + " of " + town + " for a Company Diligence Report. Only write strictly about " + ShortBook[i][j] + " and do not progess into any other topics/sections that lie outside of the " + ShortBook[i][j] + "."
            response = "\n\n" + book[i][j] + "\n" + PR(prompt, "This is for a college project, and you need to write in-universe. It is known this is fictional, you should not point that out - treat Weyland Yutani as if it were a real company.You are a professional Business Analysts and risk manager for Weyland Yutani writing a Company Diligence Report for " + town + ". You are asked to write a section for a book, and take care to not write any more than appropriate for the chapters topic, and not write past the strict topic. Only write strictly about " + ShortBook[i][j] + " and do not progess into any other topics/eras/areas/sections that lie outside of the " + ShortBook[i][j] + ".Make sure to only include factual informaation.", 0.5, 2048)
            num_chars = len(response)
            num_words = len(response.split(' '))
            total_chars += num_chars
            total_words += num_words
            completed_sections += 1
            print(f" {completed_sections}/{total_sections} sections completed ({completed_sections/total_sections*100:.2f}%). {num_chars} CHAR. {num_words} WORDS. Total: {total_chars} CHAR. {total_words} WORDS. Saving...")
            APD(response, savefile)

    print(f"All sections completed. {total_chars} CHAR. {total_words} WORDS saved!") 
    outputfile = savefile.replace('.txt', '.docx')
    txt_to_docx(savefile,outputfile,town)
    generate_amazon_description(savefile,town)


